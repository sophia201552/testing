from docx import Document
from docx.shared import Inches, Pt
import base64
import os
from beopWeb.mod_common.Utils import Utils


class WordReport:
    report = None
    document = None

    def __init__(self, report, name):
        self.report = report
        self.name = name
        self.document = Document()

    def table_cell_style(self, cell, text):
        text_formatted = cell.paragraphs[0].add_run(text)
        text_formatted.font.size = Pt(6.5)

    def save(self, path):
        if not self.report:
            return None

        for chapter in self.report:  # 报表章节
            chapter_title = chapter.get('title')
            if chapter_title:
                self.document.add_heading(chapter_title.strip(), level=1)
            unit = chapter.get('unit')
            for unit_item in unit:  # 报表章节单元
                if unit_item.get('title'):
                    self.document.add_heading(unit_item.get('title').strip(), level=2)
                content = unit_item.get('content')
                for content_item in content:  # 内容
                    if content_item.get('type') == 'summary':  # summary
                        self.document.add_paragraph(content_item.get('value').strip(), style='Body Text')
                    elif content_item.get('type') == 'chart':  # 图表
                        chart_id = content_item.get('value')
                        with open(Utils.REPORT_UPLOAD_IMAGE_FOLDER + chart_id + '.png', 'rb') as fp:
                            self.document.add_picture(fp, width=Inches(6.4527559))
                        try:
                            os.remove(Utils.REPORT_UPLOAD_IMAGE_FOLDER + chart_id)
                        except Exception as e:
                            pass

                    elif content_item.get('type') == 'table':  # table
                        header = content_item.get('header')
                        rows_data = content_item.get('rows')
                        rows_num = len(rows_data) + 1
                        columns_num = len(header)

                        table = self.document.add_table(rows=rows_num, cols=columns_num, style='Table Grid')
                        table_cells = table._cells

                        for i in range(rows_num):
                            if i == 0:  # table header
                                for h_index, header_item in enumerate(header):
                                    self.table_cell_style(table.rows[0].cells[h_index], str(header_item))
                                continue
                            row_cells = table_cells[i * columns_num:(i + 1) * columns_num]
                            for index, td in enumerate(rows_data[i - 1]):
                                self.table_cell_style(row_cells[index], str(td))
        try:
            os.remove(path)
        except Exception as e:
            pass
        with open(path, 'wb') as fh:
            self.document.save(fh)
